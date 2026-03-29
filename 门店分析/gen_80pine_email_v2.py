from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Subject
p = doc.add_paragraph()
run = p.add_run("Subject: ")
run.bold = True
run.font.size = Pt(11)
p.add_run("Re: Luckin Coffee(100 Maiden Lane) Partnership with 80 Pine St")

doc.add_paragraph("")

# Body
doc.add_paragraph("Hi Elky,")

doc.add_paragraph(
    "Thank you for your patience and for keeping in touch! "
    "Congratulations on the 40+ move-ins at Pearl & Pine \u2014 that's exciting news!"
)

doc.add_paragraph(
    "My name is Xiaoxiao, and I'll be your point of contact for this partnership going forward. "
    "Enkai has passed along your conversation, and I'm happy to get things moving."
)

doc.add_paragraph(
    "We'd love for the message below to be shared with all residents as the official announcement:"
)

# Divider
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("\u2500" * 60)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# Announcement subject
p = doc.add_paragraph()
run = p.add_run("\u3010Subject: 30% OFF Luckin Coffee Coupons \u2013 Exclusive for Pearl & Pine Residents\u3011")
run.bold = True

doc.add_paragraph("")
doc.add_paragraph("Dear Residents,")

doc.add_paragraph(
    "We are excited to share that 80 Pine St has partnered with Luckin Coffee "
    "to bring you an exclusive monthly coffee perk!"
)

# Here's what you'll get
p = doc.add_paragraph()
run = p.add_run("Here\u2019s what you\u2019ll get:")
run.bold = True

# Bullet points
bullets_data = [
    [("2 coupons", True), (" for ", False), ("30% OFF", True), (" any Luckin Coffee drinks", False)],
    [("Each resident can claim ", False), ("once per month", True)],
    [("Coupons are valid for ", False), ("30 days", True), (" after redemption", False)],
    [("Coupons refresh on the ", False), ("1st day of every month", True)],
]

for parts in bullets_data:
    p = doc.add_paragraph(style="List Bullet")
    for text, bold in parts:
        run = p.add_run(text)
        if bold:
            run.bold = True

# How to claim
p = doc.add_paragraph()
run = p.add_run("How to claim your coupons:")
run.bold = True

steps = [
    ("Download the ", False, "Luckin Coffee App", True, " from the App Store (or scan the QR code below).", False),
    ("Log in, go to ", False, "Account \u2192 My Coupons", True, ", tap ", False, "Promo Code", True, " in the top right corner, enter the code ", False, "80pine", True, ", and tap ", False, "Redeem", True, ".", False),
    ("Enjoy your exclusive resident coffee discount!", False),
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run(f"{i}. ")
    j = 0
    while j < len(step):
        text = step[j]
        bold = step[j + 1]
        run = p.add_run(text)
        if bold:
            run.bold = True
        j += 2

# QR code placeholder
p = doc.add_paragraph()
run = p.add_run("Luckin Coffee APP Download Code:")
run.bold = True
doc.add_paragraph("(QR code attached)")

# Nearest store
p = doc.add_paragraph()
run = p.add_run("Nearest Luckin Coffee Store:")
run.bold = True
p = doc.add_paragraph()
run = p.add_run("100 Maiden Lane, New York, NY 10007")
run.bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "This program will run for the next three months, so don\u2019t miss your chance "
    "each month to claim your coupons."
)

doc.add_paragraph("Enjoy your coffee!")

# Divider
p = doc.add_paragraph()
run = p.add_run("\u2500" * 60)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# Pricing note - updated formal version
doc.add_paragraph(
    "Please kindly note that the discount has been adjusted from 50% OFF to 30% OFF. "
    "The previous rate was part of our 2025 promotional pricing. Following our annual "
    "partnership review, the standard discount for 2026 has been updated to 30% OFF "
    "across all community partnerships. We remain committed to delivering great value "
    "to your residents through this program."
)

doc.add_paragraph(
    "Thanks again for coordinating this with us \u2014 we really appreciate your help!"
)

doc.add_paragraph("")
doc.add_paragraph("Best regards,")
doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("Xiaoxiao Li")
run.bold = True

doc.add_paragraph("Email: xiaoxiao.li@lkcoffee.com")
doc.add_paragraph("Website: www.luckincoffee.us")

output = "/Users/xiaoxiao/Downloads/80Pine_Partnership_Email_v2.docx"
doc.save(output)
print(f"\u5df2\u4fdd\u5b58: {output}")
