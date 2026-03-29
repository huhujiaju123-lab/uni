from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Subject
p = doc.add_paragraph()
run = p.add_run("Subject: ")
run.bold = True
run.font.size = Pt(11)
p.add_run("Re: Luckin Coffee(100 Maiden Lane) Partnership with 80 Pine St")

doc.add_paragraph("")

# Body
doc.add_paragraph("Hi Elky,")

doc.add_paragraph(
    "Thank you for your patience and for keeping in touch! "
    "Congratulations on the 40+ move-ins at Pearl & Pine — that's exciting news!"
)

doc.add_paragraph(
    "My name is Xiaoxiao, and I'll be your point of contact for this partnership going forward. "
    "Enkai has passed along your conversation, and I'm happy to get things moving."
)

doc.add_paragraph(
    "We'd love for the message below to be shared with all residents as the official announcement:"
)

# Divider
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("─" * 60)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# Announcement subject
p = doc.add_paragraph()
run = p.add_run("【Subject: 30% OFF Luckin Coffee Coupons – Exclusive for Pearl & Pine Residents】")
run.bold = True

doc.add_paragraph("")
doc.add_paragraph("Dear Residents,")

doc.add_paragraph(
    "We are excited to share that 80 Pine St has partnered with Luckin Coffee "
    "to bring you an exclusive monthly coffee perk!"
)

# Here's what you'll get
p = doc.add_paragraph()
run = p.add_run("Here's what you'll get:")
run.bold = True

# Bullet points
bullets = [
    ("2 coupons", " for ", "30% OFF", " any Luckin Coffee drinks"),
    (None, "Coupons refresh on the ", "1st day of every month", None),
    (None, "Open to all residents, with a monthly limit of 100 coupon bundles — Redeem yours while supplies last!", None, None),
]

for b in bullets:
    p = doc.add_paragraph(style="List Bullet")
    if b[0]:
        run = p.add_run(b[0])
        run.bold = True
        p.add_run(b[1])
        run = p.add_run(b[2])
        run.bold = True
        if b[3]:
            p.add_run(b[3])
    else:
        p.add_run(b[1])
        if b[2]:
            run = p.add_run(b[2])
            run.bold = True

# How to claim
p = doc.add_paragraph()
run = p.add_run("How to claim your coupons:")
run.bold = True

steps = [
    "Download the Luckin Coffee App from the App Store (or scan the QR code below).",
    "Log in, go to Account → My Coupons, tap Promo Code in the top right corner, enter the code 80pine, and tap Redeem.",
    "Enjoy your exclusive resident coffee discount!",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(f"{i}. ")
    if i == 2:
        # Bold the code
        parts = step.split("80pine")
        p.add_run(parts[0])
        run = p.add_run("80pine")
        run.bold = True
        p.add_run(parts[1])
    else:
        p.add_run(step)

# QR code placeholder
p = doc.add_paragraph()
run = p.add_run("Luckin Coffee APP Download Code:")
run.bold = True
doc.add_paragraph("(QR code attached)")

# Nearest store
p = doc.add_paragraph()
run = p.add_run("Nearest Luckin Coffee Store:")
run.bold = True
doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("100 Maiden Lane, New York, NY 10007")
run.bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "This program will run for the next three months, so don't miss your chance "
    "each month to claim your coupons."
)

doc.add_paragraph("Enjoy your coffee!")

# Divider
p = doc.add_paragraph()
run = p.add_run("─" * 60)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# Pricing note
doc.add_paragraph(
    "A quick note on pricing — as we've expanded this partnership program to more communities "
    "across the city, we've standardized the offer at 30% OFF to keep the program sustainable "
    "long-term. This ensures we can continue providing exclusive perks to your residents on an ongoing basis."
)

doc.add_paragraph(
    "Thanks again for coordinating this with us — we really appreciate your help!"
)

doc.add_paragraph("")
doc.add_paragraph("Best regards,")
doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("Xiaoxiao Li")
run.bold = True

doc.add_paragraph("Email: xiaoxiao.li@lkcoffee.com")
doc.add_paragraph("Website: www.luckincoffee.us")

output = "/Users/xiaoxiao/Downloads/80Pine_Partnership_Email.docx"
doc.save(output)
print(f"已保存: {output}")
